"""生成 docs/发展规划_agent-culture.docx（全部用单引号字符串避免双引号嵌套问题）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, pathlib

BASE = pathlib.Path(__file__).parent.parent
OUT = BASE / 'docs' / '发展规划_agent-culture.docx'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)


def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x7A)
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)


def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)


def h3(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)


def body(t):
    p = doc.add_paragraph(t); p.paragraph_format.space_after = Pt(4)
    for r in p.runs: r.font.size = Pt(11)


def bullet(t, level=0):
    p = doc.add_paragraph(style='List Bullet'); p.text = t
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(11)


def tip(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)


def code(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.left_indent = Cm(0.5)


# 封面
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Agent Culture'); r.bold = True; r.font.size = Pt(30)
r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x7A)
t.paragraph_format.space_before = Pt(60); t.paragraph_format.space_after = Pt(8)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run('发展规划与软著申请完善方案')
r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)
s.paragraph_format.space_after = Pt(6)

d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_str = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
r3 = d.add_run('版本：V1.0    日期：' + date_str)
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
d.paragraph_format.space_after = Pt(80)

doc.add_page_break()

# 一、现状
h1('一、当前版本功能总结（V1.2.1）')
body('Agent Culture 是一个面向海外市场（初期聚焦非洲地区）的 AI 视频内容创作智能体系统。截至 V1.2.1，系统已具备以下核心能力：')

h2('1.1 已完成的功能模块')
bullet('智能对话（Chat）：多轮对话 + 市场文化规则注入 + 素材引用上下文（material_ids）')
bullet('内容洞察（Insight）：URL 抓取 / 粘贴正文 → DeepSeek 文化适配分析 → 二次结构化提取（extract）→ 一键入库')
bullet('用户素材库（Materials）：个人入库 / 系统规则 / 搜索 / 市场过滤 / admin 跨用户视图（V1.2.0 新增）')
bullet('任务管理（Jobs）：视频脚本生成 + SiliconFlow Wan 视频异步任务全生命周期追踪 + 实时统计')
bullet('用户系统：JWT 登录 + 自助注册 + 角色管理（admin/user）+ 审计日志 + CSV 导出 + 管理员用户管理')
bullet('安全防护：SSRF 防护（URL 安全校验）+ 全局限流（slowapi）+ bcrypt + HS256 JWT')
bullet('多市场支持：AFRICA / US / EU / DEFAULT 四套文化规则，JSON 数据驱动无需改代码')
bullet('启动脚本：启动.bat（Windows 一键、参数化端口/热重载/浏览器）+ PowerShell quick_start.ps1')

h2('1.2 技术栈')
bullet('后端：Python 3.11 + FastAPI + SQLite（task_store.db / auth.db）')
bullet('前端：原生 HTML/CSS/JavaScript，自研 Markdown 渲染器，零构建依赖')
bullet('AI 接口：DeepSeek（文本/脚本/结构化提取）+ SiliconFlow Wan2.2-T2V（视频生成）')
bullet('部署：本地 Windows（启动.bat）/ 阿里云 ECS Linux（systemd + Nginx 反代）')

# 二、软著评估
h1('二、软著申请准备度评估')

h2('2.1 当前满足项')
bullet('[√] 系统功能完整可用：从用户注册到内容生成到素材管理形成完整闭环')
bullet('[√] 原创性充分：文化规则引擎、三步走洞察流水线、对话素材注入机制均为独创设计')
bullet('[√] 源代码可读：纯 Python + 原生 JS，无编译产物，满足"可读源代码"提交要求')
bullet('[√] 前后端分离：wangye/ 前端 + src/ 后端，结构清晰易于审核')
bullet('[√] 有一定代码量：后端约 2000+ 行 Python，前端约 2500+ 行 HTML/CSS/JS')

h2('2.2 当前不足项（需完善）')
h3('① 配套文档不足（软著必备）')
bullet('缺少《用户操作手册》：描述各功能使用步骤的图文说明文档（评审重点之一）')
bullet('缺少《软件需求规格说明书（SRS）》：功能点描述、输入输出规格等')
bullet('README / GUIDE 虽已有，但不符合软著申请要求的正式格式')

h3('② 源代码量偏少')
bullet('软著对代码量无硬性要求，但通常建议核心逻辑代码不少于 30 页（A4 10pt 字体）')
bullet('当前约可提交 40-50 页，处于基准线，建议通过新功能将代码量提升至 80+ 页')

h3('③ 功能独特性可进一步强化')
bullet('目前"视频生成"依赖 SiliconFlow 第三方，独创性主要体现在文化适配层')
bullet('建议新增更多独创算法/流程（如文化适配评分、内容日历排期、AI 建议反馈等）提升区分度')

h3('④ 界面截图材料')
bullet('申请需提交 5 张以上代表性界面截图（JPG/PNG，每张需标注功能名称）')
bullet('建议制作登录页、智能对话、内容洞察、素材库、任务管理五张标准截图')

# 三、规划
h1('三、功能发展规划路线图')
body('以下规划按优先级和开发周期分为三个阶段，每个阶段完成后作为新版本节点，同时积累软著申请素材。')

h2('3.1 近期计划（V1.3）— 打磨核心体验')
tip('预估周期：1-2 周 | 目标：提升使用流畅度、补充独创功能点')

h3('① WebSocket 实时通知（替代轮询）')
bullet('当前任务管理页采用 6s 轮询刷新，体验粗糙')
bullet('改为 FastAPI WebSocket 推送：任务状态变更时主动通知前端')
bullet('技术要点：asyncio + starlette WebSocket manager，前端 reconnect 逻辑')

h3('② 文化适配评分模块')
bullet('在"内容洞察"页新增"适配度评分"子功能')
bullet('DeepSeek 对输入文本按 5 个维度（语气/禁忌/本地化/可信度/共鸣度）各给 1-10 分')
bullet('以雷达图（前端 Canvas / SVG）可视化展示，增强独创性和视觉辨识度')
bullet('是本项目区别于普通 ChatGPT 套壳的核心差异点之一，软著申请时重点体现')

h3('③ 素材标签体系完善')
bullet('当前标签来自 AI 结构化提取，用户无法自定义编辑')
bullet('在素材详情弹窗中增加"编辑标签"功能（行内 input + Enter 确认）')
bullet('后端 materials 表增加 user_tags 字段，与 AI 提取的 tags 分开存储')

h3('④ 市场规则扩展')
bullet('新增 SEA（东南亚）市场：马来语/印尼语 / 伊斯兰饮食禁忌 / 家庭价值观语气')
bullet('新增 LATAM（拉丁美洲）市场：西班牙语 / 热情表达风格 / 足球/节日文化引用')
bullet('JSON 数据驱动，仅需新增 data/culture/sea.json 和 latam.json')

h2('3.2 中期计划（V1.4 - V1.5）— 扩展功能深度')
tip('预估周期：2-4 周 | 目标：丰富功能层次、扩大代码量、强化独创算法')

h3('① 内容日历（V1.4）')
bullet('新增"日历视图"Tab，以月历形式展示创建的任务/素材的时间分布')
bullet('支持拖拽重排、任务状态色块区分（success=绿 / running=蓝 / failed=红）')
bullet('前端用原生 CSS Grid 实现，不引入第三方日历库（保持零依赖风格）')

h3('② 素材版本历史（V1.4）')
bullet('每次 PUT 更新素材时写入 material_history 表（id, material_id, content_snapshot, modified_at）')
bullet('素材详情页"历史版本"按钮展开历史列表，可一键回滚')

h3('③ 批量操作与数据导出（V1.5）')
bullet('素材库支持多选（checkbox）+ 批量删除 / 批量导出（JSON/CSV）')
bullet('任务管理支持批量导出任务报告（后端生成 CSV）')
bullet('数据导出功能体现"数据自主权"，有助于软著申请中体现系统完整性')

h3('④ AI 内容反馈闭环（V1.5）')
bullet('对话界面 AI 回复下增加"有用 / 无用"快捷反馈按钮')
bullet('反馈数据写入 feedback 表（message_hash, rating, market, timestamp）')
bullet('admin 页面展示聚合反馈统计图（各市场满意度趋势），增加系统数据分析维度')

h2('3.3 长期规划（V2.0）— 架构升级')
tip('预估周期：4-8 周 | 目标：形成完整产品，大幅提升代码复杂度与独创性')

h3('① 协作工作区')
bullet('新增"团队素材池"概念：admin 可创建工作区，邀请成员，素材可设为"团队可见"')
bullet('权限体系扩展为 admin / editor / viewer 三级')

h3('② 品牌音调库（Brand Voice）')
bullet('用户可录入品牌关键词、禁用词、偏好句式，保存为"品牌规则包"')
bullet('在脚本生成和内容洞察时可选择附加品牌规则包，叠加到文化规则之上')
bullet('是 V2.0 最具原创性的功能，与文化规则引擎形成双引擎注入架构')

h3('③ 移动端 PWA 适配')
bullet('为 wangye/ 前端补充 manifest.json + Service Worker，实现 PWA 安装')
bullet('优化移动端 CSS 布局（折叠侧边栏 / 触摸友好按钮尺寸）')

h3('④ 本地化部署支持')
bullet('新增 Ollama 本地模型适配层（与 DeepSeek API 相同接口），支持离线运行')
bullet('对 providers.py 抽象为统一 LLMProvider 接口，不同模型只需实现接口')

# 四、申请清单
h1('四、软著申请行动清单')

h2('4.1 代码材料准备（申请时提交）')
table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '材料'; hdr[1].text = '格式/要求'; hdr[2].text = '状态'
for cell in hdr:
    for para in cell.paragraphs:
        for r in para.runs:
            r.bold = True; r.font.size = Pt(10)

rows_data = [
    ('源代码（前端）', 'wangye/*.html+*.css+*.js，打印前 30 页 + 后 30 页', '可提交'),
    ('源代码（后端）', 'src/**/*.py，打印前 30 页 + 后 30 页', '可提交'),
    ('用户操作手册', 'A4，不少于 6 页，含功能截图', '待编写'),
    ('界面截图', '5 张以上，JPG，标注功能名', '待截图'),
    ('软件说明书（SRS）', '功能规格描述，建议 5-10 页', '待编写'),
]
for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            for r in para.runs: r.font.size = Pt(10)

doc.add_paragraph()

h2('4.2 版本节点建议')
bullet('V1.3 完成后：申请时机基本成熟（含评分模块，独创性充足）')
bullet('V1.4 完成后：代码量约 80-100 页，最佳申请窗口')
bullet('申请前必须完成：用户手册 + SRS + 5 张界面截图')
tip('注意：软著申请保护的是"表达形式"（代码 + 界面），不保护算法思想。代码可读性和功能完整性比功能数量更重要。')

h2('4.3 独创性亮点梳理（供填表用）')
bullet('文化规则引擎：JSON 数据驱动 → LLM Prompt 注入，无需改代码扩展新市场（src/services/culture.py）')
bullet('三步走洞察流水线：URL 抓取 → 文化适配分析 → 二次 json_object 结构化提取 → 入库（src/api/routes/content.py）')
bullet('对话素材注入机制：用户选定素材 → 后端拼 system prompt 上下文 → LLM 引用事实作答（src/api/routes/chat.py）')
bullet('SSRF 安全防护：URL 请求前主动解析 IP，阻断私网/本机访问（src/services/url_safety.py）')
bullet('零构建前端：自研 renderMarkdown + 自研组件体系，纯原生 JS（wangye/script.js）')
bullet('双模型异步流水线：DeepSeek 同步 + SiliconFlow Wan2.2 异步轮询（src/services/pipeline.py）')

# 五、风险
h1('五、风险与注意事项')
h2('5.1 技术风险')
bullet('DeepSeek / SiliconFlow API 变更：已通过 providers.py 抽象隔离，更换模型只需修改此文件')
bullet('SQLite 并发写入：当前单机场景安全；若多进程 workers > 1 需升级 PostgreSQL')
bullet('PS5.1 编码问题：已通过 ASCII 化 ps1 文件解决，后续不向 .ps1 中写中文')

h2('5.2 软著申请风险')
bullet('第三方库版权：FastAPI / httpx / slowapi 等均为 MIT 许可证，无版权风险')
bullet('AI 生成内容版权：本系统生成内容的版权归用户所有（需在用户手册中声明）')
bullet('接口一致性：v1 路径稳定，不做破坏性变更，保证申请时代码与运行版本一致')

# 附录
h1('附录：目录结构（V1.2.1 快照）')
for line in [
    'agent-culture/',
    '|-- 启动.bat                   # Windows 一键启动（--port / --no-browser / --no-reload）',
    '|-- requirements.txt',
    '|-- data/culture/              # 市场文化规则 JSON（africa / us / eu / default）',
    '|-- docs/',
    '|   |-- GUIDE.md               # 部署与操作指南',
    '|   `-- ALIYUN_DEPLOY.md',
    '|-- scripts/',
    '|   `-- quick_start.ps1        # PowerShell 启动脚本',
    '|-- src/',
    '|   |-- main.py                # FastAPI 入口 + 审计中间件 + 限流',
    '|   |-- api/routes/            # auth / chat / content / culture / jobs',
    '|   |-- core/                  # settings / security / limiter',
    '|   `-- services/              # culture / pipeline / providers / auth_store / task_store / url_safety',
    '|-- wangye/                    # index.html / login.html / script.js / styles.css',
    '`-- storage/                   # auth.db / task_store.db',
]:
    code(line)

doc.save(str(OUT))
print('Saved:', OUT)
